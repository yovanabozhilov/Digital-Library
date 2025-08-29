from django.conf import settings
from django.shortcuts import render, redirect, get_object_or_404
from django.views.decorators.csrf import csrf_exempt
from django.views.decorators.http import require_POST
import json
import requests
from django.http import HttpResponse, FileResponse, JsonResponse
from django.contrib.auth import login, logout
from django.contrib.auth.decorators import login_required
from django.contrib.auth.forms import AuthenticationForm
from django.contrib.auth.models import User
from django.contrib import messages
from django.template.loader import render_to_string
from xhtml2pdf import pisa
from docx import Document
from bs4 import BeautifulSoup
from django.db.models import Sum, Max
from datetime import timedelta
from django.utils import timezone
from django.utils.timezone import localtime, localdate
from dateutil.relativedelta import relativedelta
from googletrans import Translator
from .forms import BookForm, EbookFileForm, ChapterForm
from .models import (Book, MyBook, Chapter, EbookFile,  SavedQuote, JournalEntry, WordLookup, Review, ReadingSession)
from django.core.files.storage import default_storage
from django.core.files.base import ContentFile
import os
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH 

def home(request):
    return render(request, 'library/welcome.html')


@login_required(login_url='login')
def main_page(request):
    user = request.user

    today = localtime().date()

    latest_sessions = (
        ReadingSession.objects.filter(user=user)
        .values('book')
        .annotate(latest_time=Max('start_time'))
        .order_by('-latest_time')[:5]
    )

    recent_books = ReadingSession.objects.filter(
        user=user,
        start_time__in=[entry['latest_time'] for entry in latest_sessions],
        book_id__in=[entry['book'] for entry in latest_sessions]
    ).order_by('-start_time')[:3]

    user_sessions = ReadingSession.objects.filter(user=user)
    stats = {
        'today_minutes': user_sessions.filter(start_time__date=today).aggregate(total=Sum('duration_minutes'))['total'] or 0,
        'today_pages': user_sessions.filter(start_time__date=today).aggregate(total=Sum('pages_read'))['total'] or 0,
    }

    read_books = Book.objects.filter(added_by=user, status='read')

    if read_books.exists():
        external_recommendations = get_google_books_recommendations(read_books)
    else:
        external_recommendations = []

    return render(request, 'library/main.html', {
        'external_recommendations': external_recommendations,
        'recent_books': recent_books,
        'stats': stats,
    })

@login_required(login_url='login')
def profile_page(request):
    def fmt_minutes(total_minutes):
        total = int(total_minutes or 0)
        h, m = divmod(total, 60)
        parts = []
        if h:
            parts.append(f"{h} час" if h == 1 else f"{h} часа")
        if m or not h:
            parts.append(f"{m} минута" if m == 1 else f"{m} минути")
        return " ".join(parts)

    today = localdate()
    week_ago = today - timedelta(days=7)
    month_ago = today - relativedelta(months=1)

    user_sessions = ReadingSession.objects.filter(user=request.user)
    raw_today = user_sessions.filter(start_time__date=today).aggregate(total=Sum('duration_minutes'))['total'] or 0
    raw_week  = user_sessions.filter(start_time__date__gte=week_ago).aggregate(total=Sum('duration_minutes'))['total'] or 0
    raw_month = user_sessions.filter(start_time__date__gte=month_ago).aggregate(total=Sum('duration_minutes'))['total'] or 0

    stats = {
        'today_minutes': fmt_minutes(raw_today),
        'week_minutes': fmt_minutes(raw_week),
        'month_minutes': fmt_minutes(raw_month),
    }

    mybooks = MyBook.objects.filter(user=request.user)
    latest_quotes = SavedQuote.objects.filter(user=request.user).order_by('-created_at')[:5]
    recent_words = WordLookup.objects.filter(user=request.user).order_by('-created_at')[:5]

    return render(request, 'library/profile.html', {
        'stats': stats,
        'mybooks': mybooks,
        'latest_quotes': latest_quotes,
        'recent_words': recent_words,
    })

def register_view(request):
    if request.method == "POST":
        username = request.POST.get("username").strip()
        email = request.POST.get("email").strip()
        password = request.POST.get("password")
        password2 = request.POST.get("password2")

        if not username or not email or not password or not password2:
            messages.error(request, "Моля, попълнете всички полета.")
        elif password != password2:
            messages.error(request, "Паролите не съвпадат!")
        elif User.objects.filter(username=username).exists():
            messages.error(request, "Потребителското име вече съществува!")
        else:
            user = User.objects.create_user(username=username, email=email, password=password)
            login(request, user)
            return redirect("main")

    return render(request, "library/register.html")

def login_view(request):
    if request.method == "POST":
        username = request.POST.get("username", "").strip()
        password = request.POST.get("password", "").strip()

        if not username or not password:
            messages.error(request, "Моля, попълнете всички полета.")
            return render(request, "library/login.html", {"form": AuthenticationForm(request)})

        form = AuthenticationForm(request, data={"username": username, "password": password})
        if form.is_valid():
            user = form.get_user()
            login(request, user)
            next_url = request.GET.get("next")
            return redirect(next_url or "main")
        else:
            messages.error(request, "Невалидни данни за вход!")
            return render(request, "library/login.html", {"form": form})

    return render(request, "library/login.html", {"form": AuthenticationForm(request)})


def logout_view(request):
    logout(request)
    return redirect("welcome")


@login_required
def read_book(request, book_id):
    book = get_object_or_404(Book, id=book_id)
    quotes = SavedQuote.objects.filter(user=request.user, book=book)
    notes = JournalEntry.objects.filter(user=request.user, book=book)
    return render(request, 'library/read_book.html', {
        'book': book,
        'quotes': quotes,
        'notes': notes,
    })


@login_required
def read_ebook(request, book_id):
    ebook = get_object_or_404(EbookFile, book_id=book_id)
    return FileResponse(ebook.file.open('rb'), content_type='application/pdf')


@login_required
def add_book(request):
    if request.method == 'POST':
        book_form = BookForm(request.POST, request.FILES)
        file_form = EbookFileForm(request.POST, request.FILES)

        if book_form.is_valid() and file_form.is_valid():
            book = book_form.save(commit=False)
            book.added_by = request.user

            if 'cover_image' in request.FILES:
                book.cover_image = request.FILES['cover_image']

            book.save()

            ebook = file_form.save(commit=False)
            ebook.book = book
            ebook.save()

            return redirect('main')
    else:
        book_form = BookForm()
        file_form = EbookFileForm()

    return render(request, 'library/add_book.html', {
        'book_form': book_form,
        'file_form': file_form,
        'MEDIA_URL': settings.MEDIA_URL
    })


@csrf_exempt
@require_POST
@login_required
def save_reading_duration(request, book_id):
    try:
        data = json.loads(request.body)
        duration = float(data.get("duration_minutes", 0))

        if duration <= 0:
            return JsonResponse({"success": False, "error": "Invalid duration"}, status=400)

        book = get_object_or_404(Book, id=book_id)
        now = timezone.now()
        start_time = now - timedelta(minutes=duration)

        ReadingSession.objects.create(
            user=request.user,
            book=book,
            start_time=start_time,
            end_time=now,
            pages_read=data.get("pages_read", 0) or 0
        )

        return JsonResponse({"success": True})
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=400)
    

GOOGLE_BOOKS_API_URL = "https://www.googleapis.com/books/v1/volumes"

def get_google_books_recommendations(read_books):
    recommendations = []
    seen_titles = set()

    existing_titles = set(Book.objects.values_list('title', flat=True))

    preferred_authors = list(set(book.author for book in read_books if book.author))
    preferred_genres = list(set(book.genre for book in read_books if book.genre))
    preferred_moods = list(set(book.mood for book in read_books if book.mood))

    queries = []

    for author in preferred_authors:
        queries.append(f"inauthor:{author}")

    for genre in preferred_genres:
        queries.append(f"subject:{genre}")

    for mood in preferred_moods:
        queries.append(mood)

    for query in queries:
        params = {
            'q': query,
            'maxResults': 3,
            'printType': 'books',
            'langRestrict': 'bg',
        }
        response = requests.get(GOOGLE_BOOKS_API_URL, params=params)

        if response.status_code == 200:
            items = response.json().get('items', [])
            for item in items:
                volume_info = item.get('volumeInfo', {})
                title = volume_info.get('title')

                if title and title not in seen_titles and title not in existing_titles:
                    recommendations.append({
                        'title': title,
                        'authors': ', '.join(volume_info.get('authors', [])),
                        'description': volume_info.get('description', '')[:300],
                        'thumbnail': volume_info.get('imageLinks', {}).get('thumbnail', ''),
                        'link': volume_info.get('infoLink', '#')
                    })
                    seen_titles.add(title)

        if len(recommendations) >= 10:
            break

    if not recommendations:
        fallback_params = {
            'q': 'books',  
            'maxResults': 10,
            'printType': 'books',
            'langRestrict': 'bg',
        }
        fallback_response = requests.get(GOOGLE_BOOKS_API_URL, params=fallback_params)
        if fallback_response.status_code == 200:
            items = fallback_response.json().get('items', [])
            for item in items:
                volume_info = item.get('volumeInfo', {})
                title = volume_info.get('title')
                if title and title not in seen_titles and title not in existing_titles:
                    recommendations.append({
                        'title': title,
                        'authors': ', '.join(volume_info.get('authors', [])),
                        'description': volume_info.get('description', '')[:300],
                        'thumbnail': volume_info.get('imageLinks', {}).get('thumbnail', ''),
                        'link': volume_info.get('infoLink', '#')
                    })
                    seen_titles.add(title)
                if len(recommendations) >= 10:
                    break

    return recommendations[:10]


@login_required
def book_detail(request, book_id):
    book = get_object_or_404(Book, pk=book_id)
    reviews = Review.objects.filter(book=book).order_by('-created_at')
    user_review = None

    if request.method == 'POST':
        if 'delete_review_id' in request.POST:
            review_to_delete = Review.objects.filter(id=request.POST['delete_review_id'], user=request.user).first()
            if review_to_delete:
                review_to_delete.delete()
            return redirect('book_detail', book_id=book.id)

        elif 'edit_review_id' in request.POST:
            review_id = request.POST['edit_review_id']
            user_review = Review.objects.filter(id=review_id, user=request.user).first()

        else:
            rating = int(request.POST.get('rating'))
            content = request.POST.get('content').strip()
            review_id = request.POST.get('review_id')

            if review_id:
                review = Review.objects.filter(id=review_id, user=request.user).first()
                if review:
                    review.rating = rating
                    review.content = content
                    review.save()
            else:
                Review.objects.create(book=book, user=request.user, rating=rating, content=content)
            return redirect('book_detail', book_id=book.id)

    genre_dict = dict(BookForm.PREDEFINED_GENRES)
    mood_dict = dict(BookForm.PREDEFINED_MOODS)

    genres = book.genre.split(',') if book.genre else []
    translated_genres = ', '.join([genre_dict.get(g.strip(), g.strip()) for g in genres])

    moods = book.mood.split(',') if book.mood else []
    translated_moods = ', '.join([mood_dict.get(m.strip(), m.strip()) for m in moods])

    return render(request, 'library/book_detail.html', {
        'book': book,
        'reviews': reviews,
        'user_review': user_review,
        'genre_display': translated_genres,
        'mood_display': translated_moods,
    })


@login_required
def my_added_books(request):
    books = Book.objects.filter(added_by=request.user)

    genre = request.GET.get('genre', '')
    mood = request.GET.get('mood', '')
    title = request.GET.get('title', '')
    author = request.GET.get('author', '')

    if genre:
        books = books.filter(genre__icontains=genre) 

    if mood:
        books = books.filter(mood__icontains=mood)   

    if title:
        books = books.filter(title__icontains=title)

    if author:
        books = books.filter(author__icontains=author)

    books = books.order_by('title')

    genres = BookForm.PREDEFINED_GENRES
    moods = BookForm.PREDEFINED_MOODS

    return render(request, 'library/my_added_books.html', {
        'books': books,
        'genres': genres,
        'moods': moods,
        'selected': {
            'genre': genre,
            'mood': mood,
            'title': title,
            'author': author
        }
    })

@login_required
def delete_book(request, book_id):
    book = get_object_or_404(Book, id=book_id, added_by=request.user)
    
    if request.method == 'POST':
        book.delete()
        return redirect('my_added_books')
    
    return redirect('book_detail', book_id=book_id)

@csrf_exempt
@require_POST
@login_required
def save_last_page(request, book_id):
    try:
        data = json.loads(request.body)
        page = int(data.get("page", 1))

        ebook = get_object_or_404(EbookFile, book_id=book_id, book__added_by=request.user)
        ebook.last_page_read = page
        ebook.save()

        return JsonResponse({"success": True, "page": page})
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=400)
    

@login_required
def create_mybook(request):
    if request.method == 'POST':
        title = request.POST.get('title')
        if title:
            book = MyBook.objects.create(user=request.user, title=title)
            return redirect('view_mybook', book_id=book.id)
    return render(request, 'mybooks/create_mybook.html')


@login_required
def view_mybook(request, book_id):
    book = get_object_or_404(MyBook, id=book_id, user=request.user)

    if request.method == 'POST':
        new_title = request.POST.get('new_title')
        if new_title:
            book.title = new_title
            book.save()
            return redirect('view_mybook', book_id=book.id)

    chapters = book.chapters.all()
    return render(request, 'mybooks/view_mybook.html', {
        'book': book,
        'chapters': chapters,
    })

@login_required
def delete_mybook(request, book_id):
    mybook = get_object_or_404(MyBook, id=book_id, user=request.user)
    mybook.delete()
    return redirect('profile')  

@login_required
def add_chapter(request, book_id, chapter_id=None):
    mybook = get_object_or_404(MyBook, pk=book_id, user=request.user)
    chapter = get_object_or_404(Chapter, pk=chapter_id, book=mybook) if chapter_id else None

    if request.method == 'POST':
        form = ChapterForm(request.POST, instance=chapter)
        if form.is_valid():
            new_chapter = form.save(commit=False)
            new_chapter.mybook = mybook
            new_chapter.save()
            return redirect('view_mybook', book_id=mybook.id)
    else:
        form = ChapterForm(instance=chapter)

    return render(request, 'mybooks/add_chapter.html', {
        'form': form,
        'book': mybook,
        'chapter': chapter
    })

@login_required
def edit_chapter(request, book_id, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id, mybook__id=book_id, mybook__user=request.user)
    mybook = chapter.mybook
    if request.method == 'POST':
        chapter.title = request.POST.get('title')
        chapter.content = request.POST.get('content')
        chapter.save()
        return redirect('view_mybook', book_id=book_id)
    return render(request, 'mybooks/add_chapter.html', {'chapter': chapter, 'book': mybook})


@login_required
def delete_chapter(request, book_id, chapter_id):
    chapter = get_object_or_404(Chapter, id=chapter_id, mybook__id=book_id, mybook__user=request.user)
    chapter.delete()
    return redirect('view_mybook', book_id=book_id)


def parse_html_to_docx_paragraph(html_content, document):
    soup = BeautifulSoup(html_content, 'html.parser')

    for element in soup.find_all(recursive=False):  
        if element.name == 'p':
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  

            for child in element.children:
                text = child.get_text(strip=True) if hasattr(child, 'get_text') else str(child).strip()
                if not text:
                    continue  

                run = p.add_run(text)
                run.font.size = Pt(12)

                if hasattr(child, 'name'):
                    if child.name in ['strong', 'b']:
                        run.bold = True
                    if child.name in ['em', 'i']:
                        run.italic = True
                    if child.name == 'u':
                        run.underline = True

        elif element.name == 'h1':
            p = document.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(element.get_text(strip=True))
            run.bold = True
            run.font.size = Pt(18)

        elif element.name == 'h2':
            p = document.add_paragraph()
            run = p.add_run(element.get_text(strip=True))
            run.bold = True
            run.font.size = Pt(16)

        elif element.name == 'br':
            document.add_paragraph()  


@login_required
def export_mybook_docx(request, mybook_id):
    mybook = get_object_or_404(MyBook, id=mybook_id, user=request.user)
    chapters = mybook.chapters.all()

    doc = Document()

    title = doc.add_heading(mybook.title, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for idx, chapter in enumerate(chapters):
        if idx > 0:
            doc.add_page_break() 
        doc.add_heading(chapter.title, level=2)
        parse_html_to_docx_paragraph(chapter.content, doc)

    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="{mybook.title}.docx"'
    doc.save(response)
    return response


@login_required
def export_mybook_pdf(request, mybook_id):
    mybook = get_object_or_404(MyBook, id=mybook_id, user=request.user)
    chapters = mybook.chapters.all()

    html = render_to_string('mybooks/mybook_pdf_template.html', {
        'mybook': mybook,
        'chapters': chapters
    })

    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="{mybook.title}.pdf"'

    pisa.CreatePDF(html, dest=response)
    return response


@csrf_exempt 
@login_required
def save_quote(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            text = data.get("text", "").strip()
            book_id = data.get("book_id")
            page = data.get("page", None)  

            if text and book_id:
                SavedQuote.objects.create(
                    user=request.user,
                    book_id=book_id,
                    quote_text=text,
                    page=page if page else None,
                    created_at=timezone.now()
                )
                return JsonResponse({"status": "ok"})
            else:
                return JsonResponse({"status": "missing fields"}, status=400)
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)

    return JsonResponse({"status": "invalid request"}, status=400)

@csrf_exempt
@login_required
def save_note(request):
    if request.method == "POST":
        data = json.loads(request.body)
        content = data.get("content", "").strip()
        book_id = data.get("book_id")
        page = data.get("page")
        if content and book_id:
            JournalEntry.objects.create(
                user=request.user,
                book_id=book_id,
                content=content,
                page=page 
            )
            return JsonResponse({"status": "ok"})
    return JsonResponse({"status": "error"}, status=400)

@csrf_exempt
@login_required
def delete_quote(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            quote_id = data.get("id")
            quote = SavedQuote.objects.get(id=quote_id, user=request.user)
            quote.delete()
            return JsonResponse({"status": "ok"})
        except SavedQuote.DoesNotExist:
            return JsonResponse({"status": "not found"}, status=404)
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)
    return JsonResponse({"status": "invalid"}, status=400)

@csrf_exempt
@login_required
def delete_note(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            note_id = data.get("id")
            note = JournalEntry.objects.get(id=note_id, user=request.user)
            note.delete()
            return JsonResponse({"status": "ok"})
        except JournalEntry.DoesNotExist:
            return JsonResponse({"status": "not found"}, status=404)
        except Exception as e:
            return JsonResponse({"status": "error", "message": str(e)}, status=500)
    return JsonResponse({"status": "invalid"}, status=400)

@login_required
def lookup_word(request):
    query = request.GET.get("q", "").strip()
    result = None
    if query:
        result = WordLookup.objects.filter(word__iexact=query).first()
    return render(request, "library/lookup_result.html", {"query": query, "result": result})

@login_required
def translate_and_define(request):
    word = request.GET.get("word", "").strip()

    if not word:
        return JsonResponse({"error": "Няма въведена дума."})

    try:
        translator = Translator()

        is_cyrillic = all('а' <= c.lower() <= 'я' for c in word if c.isalpha())

        if is_cyrillic:
            translation = translator.translate(word, src='bg', dest='en')
            lookup_word = translation.text
            display_translation = word
        else:
            lookup_word = word
            translation = translator.translate(word, src='en', dest='bg')
            display_translation = translation.text

        definitions = {}
        dict_response = requests.get(f"https://api.dictionaryapi.dev/api/v2/entries/en/{lookup_word}")

        if dict_response.status_code == 200:
            dict_data = dict_response.json()
            meanings = dict_data[0].get("meanings", [])

            for meaning in meanings:
                part = meaning.get("partOfSpeech", "").strip().lower()
                for d in meaning.get("definitions", []):
                    eng_def = d.get("definition", "")
                    if eng_def:
                        bul_def = translator.translate(eng_def, src='en', dest='bg').text
                        definitions.setdefault(part, []).append(bul_def)

        if request.user.is_authenticated:
            WordLookup.objects.update_or_create(
                user=request.user,
                word=word,
                defaults={
                    'definition': display_translation,
                    'created_at': timezone.now()
                }
            )

        return JsonResponse({
            "original": word,
            "translated": display_translation,
            "definitions": definitions,
            "message": "Няма намерена дефиниция." if not definitions else ""
        })

    except Exception as e:
        return JsonResponse({"error": f"Грешка: {str(e)}"})


@csrf_exempt
def upload_image(request):
    if request.method == 'POST' and (request.FILES.get('upload') or request.FILES.get('file')):
        image = request.FILES.get('upload') or request.FILES.get('file')
        path = default_storage.save('uploads/' + image.name, ContentFile(image.read()))
        image_url = default_storage.url(path)
        return JsonResponse({"url": image_url})
    return JsonResponse({'error': 'Invalid request'}, status=400)